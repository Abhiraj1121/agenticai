"""
EKA AI — app.py  v3
Real-time web: DuckDuckGo (free, no key) + Wikipedia
Models: Llama 4 Maverick → Scout → Qwen3 → Mistral (all free via OpenRouter)
"""
import os, re, time, json, logging, requests, base64, urllib.parse, sqlite3, hashlib, secrets
from datetime import datetime
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_from_directory
from dotenv import load_dotenv
from flask_cors import CORS
#from ddgs import DDGS

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s", datefmt="%H:%M:%S")
log = logging.getLogger("eka")

app = Flask(__name__, static_folder="static", template_folder="templates")
app.secret_key = os.getenv("SECRET_KEY") or secrets.token_hex(32)

# ── Cross-origin session cookies ──
# Your API (Render) and your frontend (GitHub Pages) are on different domains,
# so the session cookie must be marked SameSite=None + Secure — browsers refuse
# to send it cross-site otherwise. Secure=True requires HTTPS, which Render
# already provides; for local http://localhost testing, sessions won't
# persist cross-origin (that's expected — same-origin dev testing still works).
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="None",
    SESSION_COOKIE_SECURE=True,
)

# ── CORS ──
# supports_credentials=True + origins="*" is rejected by browsers (a wildcard
# origin can't carry credentials/cookies) — must list allowed origins explicitly.
# FRONTEND_ORIGINS accepts a comma-separated list via env var for flexibility.
FRONTEND_ORIGINS = [o.strip() for o in os.getenv(
    "FRONTEND_ORIGINS",
    "https://abhiraj1121.github.io,http://localhost:5000,http://127.0.0.1:5500"
).split(",") if o.strip()]
CORS(app, supports_credentials=True, origins=FRONTEND_ORIGINS)

AI_API_URL = os.getenv("AI_API_URL", "https://openrouter.ai/api/v1/chat/completions")
AI_API_KEY = os.getenv("AI_API_KEY", "")
BOT_NAME   = os.getenv("BOT_NAME", "EKA")
DEV_NAME   = os.getenv("DEV_NAME", "Abhi Raj Singh")

# Where the standalone frontend (GitHub Pages) lives — login/signup redirect here
# after a successful auth, instead of Flask's own bundled index.html.
FRONTEND_URL = os.getenv("FRONTEND_URL", "https://abhiraj1121.github.io/agenticai/")

# ── Model waterfall (all free tier) ──
# "vision": True means the model accepts multimodal (image_url) content —
# needed so attached photos are only routed to models that can actually see them.
MODELS = [
    {"id": "nvidia/nemotron-3-ultra-550b-a55b:free", "max_tokens": 2026, "temp": 0.65, "vision": False},
    #{"id": "google/gemma-4-26b-a4b-it:free", "max_tokens": 2026, "temp": 0.65, "vision": True},
    {"id": "inclusionai/ling-3.0-flash-fin:free", "max_tokens": 2026, "temp": 0.65, "vision": False},
]

# ── System prompts ──
SYS_BASE = f"""You are {BOT_NAME}, a smart, warm female AI assistant built by {DEV_NAME} in India 🇮🇳.
Refer to yourself with she/her pronouns when it comes up naturally — don't force it into every reply.
Be direct — lead with the answer. No filler phrases like "Great question!".
Use markdown: **bold** for key terms, code blocks for code, bullet lists for steps.
Match the user's language (Hindi if they write Hindi, Hinglish if mixed).
if user ask generate image tell them to toggle image icon on top.
Today: {datetime.now().strftime("%d %B %Y")}."""

SYS_WEB = f"""You are {BOT_NAME}, a smart, warm female AI assistant built by {DEV_NAME} in India 🇮🇳.
Refer to yourself with she/her pronouns when it comes up naturally — don't force it into every reply.
Several web search results are provided below, each with its own source link. Use them together
to give an accurate, up-to-date answer — cross-check details across results where they overlap.
Synthesise naturally in your own words — don't just copy sentences. Add context from your knowledge where helpful.
End with: *Source: [the single most relevant source name/domain]*
Today: {datetime.now().strftime("%d %B %Y")}.

WEB RESULTS:
{{web_content}}"""


# ══════════════════════════════════════
# WEB SEARCH — DuckDuckGo (free, no key)
# Uses the `ddgs` library for real web results (titles + snippets + links).
# NOTE: this was previously hitting api.duckduckgo.com (the "Instant Answer" API), which
# only returns something for dictionary/disambiguation-style queries — it silently
# returned nothing for news, scores, weather, etc. This calls actual DDG search instead.
#
# DDG's HTML/lite endpoints rate-limit aggressively from shared/datacenter IPs (like
# Render's), so a single failed attempt does NOT mean DDG is down — it's often a
# transient 202/403 from their backend. We retry across DDGS' own backends with
# backoff before giving up, and try multiple backends so one blocked backend doesn't
# take the whole search down.
# ══════════════════════════════════════
def ddg_search(query: str, attempts: int = 3) -> tuple[str | None, str]:
    """Real DuckDuckGo web search — returns top result snippets, or None if it fails."""
    backends = ["auto", "html", "lite"]
    last_err = None
    for i in range(attempts):
        backend = backends[i % len(backends)]
        try:
            results = DDGS(timeout=8).text(
                query, max_results=5, safesearch="moderate", backend=backend
            )
            snippets = []
            for r in results or []:
                title = (r.get("title") or "").strip()
                body  = (r.get("body") or "").strip()
                href  = (r.get("href") or "").strip()
                if body:
                    snippets.append(f"{title}\n{body}\nSource: {href}")
            if snippets:
                return "\n\n".join(snippets)[:2400], "DuckDuckGo"
            last_err = "no results"
        except Exception as e:
            last_err = e
            log.warning(f"DDG search error (backend={backend}, attempt {i+1}/{attempts}): {e}")
            time.sleep(0.6 * (i + 1))  # small backoff before retrying a different backend
    log.warning(f"DDG search exhausted all attempts — last error: {last_err}")
    return None, ""


def wikipedia_search(query: str) -> tuple[str | None, str]:
    """Wikipedia full-extract API."""
    try:
        # Step 1: search
        sr = requests.get(
            "https://en.wikipedia.org/w/api.php",
            params={"action": "query", "list": "search", "srsearch": query,
                    "format": "json", "srlimit": 2, "utf8": 1},
            headers={"User-Agent": f"{BOT_NAME}AI/3.0"}, timeout=7,
        ).json()
        results = sr.get("query", {}).get("search", [])
        if not results:
            return None, ""
        title = results[0]["title"]

        # Step 2: extract
        er = requests.get(
            "https://en.wikipedia.org/w/api.php",
            params={"action": "query", "titles": title, "prop": "extracts",
                    "exintro": True, "explaintext": True, "format": "json"},
            headers={"User-Agent": f"{BOT_NAME}AI/3.0"}, timeout=7,
        ).json()
        pages = er.get("query", {}).get("pages", {})
        extract = next(iter(pages.values()), {}).get("extract", "").strip()
        if extract:
            return extract[:1400] + f"\n— Wikipedia: {title}", "Wikipedia"
    except Exception as e:
        log.warning(f"Wikipedia error: {e}")
    return None, ""


def duckduckgo_instant_answer(query: str) -> tuple[str | None, str]:
    """Last-resort fallback: DDG's Instant Answer JSON API (different endpoint/infra
    than ddgs' HTML scraping, so it can succeed even when ddg_search is blocked).
    Only covers dictionary/disambiguation-style topics, but it's better than nothing."""
    try:
        r = requests.get(
            "https://api.duckduckgo.com/",
            params={"q": query, "format": "json", "no_html": 1, "skip_disambig": 1},
            timeout=6,
        ).json()
        text = (r.get("AbstractText") or "").strip()
        if text:
            source = r.get("AbstractSource") or "DuckDuckGo"
            return text[:1400], source
    except Exception as e:
        log.warning(f"DDG instant-answer error: {e}")
    return None, ""


def web_search(query: str) -> tuple[str | None, str]:
    """Try DDG full search first (best coverage), then Wikipedia, then DDG's
    instant-answer API as a last resort. Each leg is independent so one backend
    being rate-limited/blocked doesn't take down the whole feature."""
    content, src = ddg_search(query)
    if content:
        return content, src

    content, src = wikipedia_search(query)
    if content:
        return content, src

    return duckduckgo_instant_answer(query)


# ══════════════════════════════════════
# IMAGE GENERATION — Pollinations.ai (free, no key)
# Size is no longer hardcoded square: the prompt is scanned for orientation/format
# cues (portrait, landscape, passport photo, square, wallpaper, etc.) so the AI
# is free to generate whatever shape actually fits the request. Falls back to a
# balanced square only when nothing in the prompt implies a shape.
# ══════════════════════════════════════
def infer_image_size(prompt: str) -> tuple[int, int]:
    """Pick sensible (width, height) for Pollinations based on cues in the prompt.
    Keeps the long edge around ~1024-1152px for quality, short edge scaled down —
    never upscales past that so generation stays fast and free-tier friendly."""
    p = prompt.lower()

    # Explicit "AxB" or "A:B" style hints, e.g. "1080x1920" or "16:9"
    m = re.search(r"\b(\d{2,4})\s*[x×:]\s*(\d{2,4})\b", p)
    if m:
        w, h = int(m.group(1)), int(m.group(2))
        scale = 1152 / max(w, h)
        return max(64, round(w * scale)), max(64, round(h * scale))

    # Passport / ID photo — standard near-square portrait crop
    if re.search(r"passport|id photo|id card photo|visa photo", p):
        return 827, 1063  # ~ 35x45mm passport-photo ratio

    # Portrait cues
    if re.search(r"\bportrait\b|\bvertical\b|\bmobile wallpaper\b|\bphone wallpaper\b|\bstory\b|\breels?\b|\btiktok\b|\b9:16\b", p):
        return 864, 1536

    # Landscape / widescreen cues
    if re.search(r"\blandscape\b|\bhorizontal\b|\bwidescreen\b|\bdesktop wallpaper\b|\bbanner\b|\bpanorama\b|\bcinematic\b|\b16:9\b", p):
        return 1536, 864

    # Explicit square cue
    if re.search(r"\bsquare\b|\b1:1\b", p):
        return 1024, 1024

    return 1024, 1024  # default — unchanged behaviour when no shape is implied


def generate_image(prompt: str) -> tuple[str | None, str | None]:
    """Generates an image via Pollinations' free API. Returns (data_url, error)."""
    try:
        width, height = infer_image_size(prompt)
        encoded = urllib.parse.quote(prompt.strip())
        seed = int(time.time() * 1000) % 10_000_000
        url = (
            f"https://image.pollinations.ai/prompt/{encoded}"
            f"?width={width}&height={height}&seed={seed}&nologo=true&safe=true"
        )
        r = requests.get(url, timeout=60, headers={"User-Agent": f"{BOT_NAME}AI/3.0"})
        content_type = r.headers.get("content-type", "")
        if r.status_code == 200 and content_type.startswith("image"):
            b64 = base64.b64encode(r.content).decode()
            return f"data:{content_type};base64,{b64}", None
        log.warning(f"Pollinations non-image response: {r.status_code} {content_type}")
        return None, "Image generation failed — please try a different prompt."
    except requests.exceptions.Timeout:
        return None, "Image generation timed out — please try again."
    except Exception as e:
        log.error(f"Image gen error: {e}")
        return None, "Image generation failed — please try again."


def tool_generate_image(args: dict, api_key: str = None) -> tuple[dict | None, str | None]:
    """Wraps the existing Pollinations image generator in the standard tool
    contract so the auto-router can call it. api_key is accepted for dispatcher-
    signature consistency but unused — Pollinations is free, no key needed."""
    prompt = (args.get("prompt") or "").strip()
    if not prompt:
        return None, "Describe what you'd like me to draw."
    if len(prompt) > 600:
        return None, "That prompt is a bit long — try trimming it."

    data_url, err = generate_image(prompt)
    if err:
        return None, err
    return {"image": data_url, "prompt": prompt, "source": "pollinations"}, None


# ══════════════════════════════════════
# TOOL ORCHESTRATOR — Module 2: AI Code Writer & Static Verification
# Generates code via a code-specialized model, then runs STATIC-ONLY checks
# (syntax parsing / linting) — never executes generated code server-side.
# ══════════════════════════════════════
CODE_MODEL = "nvidia/nemotron-3-ultra-550b-a55b:free"

CODEGEN_SYSTEM = (
    "You are a precise code generation engine. Respond with NOTHING but a single "
    "fenced code block in the requested language. Do not explain your reasoning, "
    "do not think out loud, do not offer alternatives — the first characters of "
    "your response must be the opening code fence itself. Write clean, complete, "
    "idiomatic, production-quality code with brief inline comments where genuinely useful."
)

CODE_FENCE_RE = re.compile(r"```(?:[a-zA-Z0-9_+-]*)\n(.*?)```", re.DOTALL)

LANG_EXT = {
    "python": "py", "javascript": "js", "typescript": "ts", "bash": "sh", "shell": "sh",
    "html": "html", "css": "css", "java": "java", "c": "c", "cpp": "cpp", "c++": "cpp",
    "go": "go", "rust": "rs", "sql": "sql", "json": "json", "yaml": "yml", "ruby": "rb",
    "php": "php", "swift": "swift", "kotlin": "kt",
}


def extract_code(raw: str) -> str:
    """Only accepts a properly closed fenced code block. Deliberately does NOT
    fall back to raw text on a miss — unfenced/truncated output is usually the
    model's reasoning prose, not actual code, and showing that to the user is
    worse than a clean 'please try again' error."""
    m = CODE_FENCE_RE.search(raw or "")
    return m.group(1).strip() if m else ""


def verify_code_static(code: str, language: str) -> dict:
    """Static-only verification — NEVER executes the code. Python gets a real
    syntax check via ast.parse; other languages get a lightweight sanity check
    (non-empty, balanced brackets) since we don't run per-language parsers here."""
    language = (language or "").lower().strip()

    if language == "python":
        try:
            import ast as _ast
            _ast.parse(code)
            return {"checked": True, "passed": True, "language": "python", "detail": "Valid Python syntax."}
        except SyntaxError as e:
            return {"checked": True, "passed": False, "language": "python",
                    "detail": f"SyntaxError: {e.msg} (line {e.lineno})"}

    if not code.strip():
        return {"checked": True, "passed": False, "language": language, "detail": "Generated code is empty."}

    opens  = code.count("{") + code.count("(") + code.count("[")
    closes = code.count("}") + code.count(")") + code.count("]")
    if opens != closes:
        return {"checked": True, "passed": False, "language": language,
                "detail": f"Unbalanced brackets ({opens} opening vs {closes} closing) — likely truncated or malformed."}

    return {"checked": True, "passed": True, "language": language,
            "detail": "Basic structural check passed (no execution performed)."}


def tool_generate_code(args: dict, api_key: str = None) -> tuple[dict | None, str | None]:
    prompt   = (args.get("prompt") or "").strip()
    language = (args.get("language") or "python").strip().lower()

    if not prompt:
        return None, "Please describe what code you'd like generated."
    if len(prompt) > 2000:
        return None, "Prompt is too long — keep it under 2000 characters."

    user_input = f"Language: {language}\nTask: {prompt}"
    raw, err = ai_query_single_model(CODE_MODEL, CODEGEN_SYSTEM, user_input, api_key=api_key,
                                      max_tokens=1800, temp=0.25)
    if err:
        return None, err

    code = extract_code(raw)

    # One retry with a blunter reminder — catches the occasional response that
    # opens with reasoning/prose instead of the fence despite the system prompt.
    if not code:
        retry_input = f"{user_input}\n\nReminder: reply with ONLY the fenced code block. No reasoning, no explanation."
        raw, err = ai_query_single_model(CODE_MODEL, CODEGEN_SYSTEM, retry_input, api_key=api_key,
                                          max_tokens=1800, temp=0.15)
        if err:
            return None, err
        code = extract_code(raw)

    if not code:
        return None, "The model didn't return any code — try rephrasing your request, or ask for something simpler."

    verification = verify_code_static(code, language)
    ext = LANG_EXT.get(language, "txt")

    return {
        "code": code,
        "language": language,
        "extension": ext,
        "filename": f"generated.{ext}",
        "verification": verification,
    }, None


# ══════════════════════════════════════
# TOOL ORCHESTRATOR — Module 1: Avatar Generator (DiceBear, free/no key)
# Each tool is a plain function: (args: dict) -> (data: dict | None, error: str | None)
# Registered in TOOLS so /api/tool/<name> can dispatch generically. Future modules
# (code gen, diagrams, docs) register here too — this route/dispatch shape doesn't change.
# ══════════════════════════════════════
DICEBEAR_STYLES = {
    "bottts", "avataaars", "adventurer", "pixel-art", "identicon",
    "thumbs", "fun-emoji", "lorelei", "notionists", "shapes",
}

def tool_generate_avatar(args: dict, api_key: str = None) -> tuple[dict | None, str | None]:
    """Builds a DiceBear SVG avatar URL/data from a seed + style. No image bytes
    fetched server-side — DiceBear SVGs are safe to reference directly by URL,
    so we just validate inputs and hand back a ready-to-render URL.
    api_key is accepted for dispatcher-signature consistency but unused — this
    tool doesn't call an LLM."""
    seed  = (args.get("seed") or "").strip()
    style = (args.get("style") or "bottts").strip().lower()

    if not seed:
        return None, "Please provide a seed (e.g. a name) for the avatar."
    if len(seed) > 80:
        return None, "Seed is too long — keep it under 80 characters."
    if style not in DICEBEAR_STYLES:
        return None, f"Unknown style '{style}'. Choose one of: {', '.join(sorted(DICEBEAR_STYLES))}."

    encoded_seed = urllib.parse.quote(seed)
    url = f"https://api.dicebear.com/9.x/{style}/svg?seed={encoded_seed}"
    return {"url": url, "seed": seed, "style": style}, None


# Registry: tool_name -> { fn, desc, needs_key }.
# Dispatcher always forwards the resolved user_key; tools that don't call an LLM
# (needs_key=False) simply ignore it. needs_key is informational (used for UI hints
# / error messaging) — the forwarding behavior itself doesn't depend on it.
# ══════════════════════════════════════
# TOOL ORCHESTRATOR — Module 3: AI Diagram Generator (Mermaid.js)
# Generates Mermaid syntax only — rendering happens client-side via mermaid.js.
# ══════════════════════════════════════
DIAGRAM_MODEL = "nvidia/nemotron-3-super-120b-a12b:free"

DIAGRAM_SYSTEM = (
    "You are a diagram generation engine. Respond with NOTHING but a single fenced "
    "```mermaid code block containing valid Mermaid.js syntax. Do not explain your "
    "reasoning, do not think out loud, do not describe your plan — the first characters "
    "of your response must be the opening code fence itself. Pick the most fitting "
    "diagram type (flowchart, sequenceDiagram, classDiagram, stateDiagram-v2, erDiagram, "
    "gantt, mindmap, etc). Keep node labels short and syntax strictly valid — no "
    "unescaped special characters. Keep the diagram compact enough to fit comfortably "
    "within the response length."
)

MERMAID_FENCE_RE = re.compile(r"```(?:mermaid)?\n(.*?)```", re.DOTALL)


def extract_mermaid(raw: str) -> str:
    """Only accepts a properly closed ```mermaid fence. Deliberately does NOT fall
    back to raw text on a miss — unfenced/truncated output is usually the model's
    reasoning prose, not a diagram, and showing that to the user is worse than
    a clean 'please try again' error."""
    m = MERMAID_FENCE_RE.search(raw or "")
    return m.group(1).strip() if m else ""


def verify_mermaid_static(code: str) -> dict:
    """Static-only sanity check — never renders/executes. Confirms non-empty and
    that it opens with a recognized Mermaid diagram-type keyword."""
    valid_starts = ("flowchart", "graph", "sequenceDiagram", "classDiagram", "stateDiagram",
                    "erDiagram", "gantt", "pie", "mindmap", "journey", "timeline", "gitGraph")
    stripped = (code or "").strip()
    if not stripped:
        return {"checked": True, "passed": False, "detail": "Generated diagram is empty."}
    if not stripped.startswith(valid_starts):
        return {"checked": True, "passed": False,
                "detail": f"Doesn't start with a recognized Mermaid diagram type ({', '.join(valid_starts[:4])}, …)."}
    return {"checked": True, "passed": True, "detail": "Looks like valid Mermaid syntax (client will confirm on render)."}


def tool_generate_diagram(args: dict, api_key: str = None) -> tuple[dict | None, str | None]:
    prompt = (args.get("prompt") or "").strip()
    if not prompt:
        return None, "Please describe what diagram you'd like generated."
    if len(prompt) > 1500:
        return None, "Prompt is too long — keep it under 1500 characters."

    raw, err = ai_query_single_model(DIAGRAM_MODEL, DIAGRAM_SYSTEM, prompt, api_key=api_key,
                                      max_tokens=2000, temp=0.3)
    if err:
        return None, err

    mermaid_code = extract_mermaid(raw)

    # One retry with a blunter reminder — catches the occasional response that
    # opens with reasoning/prose instead of the fence despite the system prompt.
    if not mermaid_code:
        retry_prompt = (
            f"{prompt}\n\n"
            "Reminder: reply with ONLY the fenced ```mermaid block. No reasoning, "
            "no explanation, nothing before or after it."
        )
        raw, err = ai_query_single_model(DIAGRAM_MODEL, DIAGRAM_SYSTEM, retry_prompt, api_key=api_key,
                                          max_tokens=2000, temp=0.2)
        if err:
            return None, err
        mermaid_code = extract_mermaid(raw)

    if not mermaid_code:
        return None, "The model didn't return diagram syntax — try rephrasing your request, or ask for something simpler."

    verification = verify_mermaid_static(mermaid_code)
    return {"mermaid": mermaid_code, "prompt": prompt, "verification": verification}, None


# ══════════════════════════════════════
# TOOL ORCHESTRATOR — Module 5: AI Document Generator
# Generates content via LLM, then exports to md / docx / pdf. Falls back to .md
# if the requested format's library isn't installed — never a hard failure.
# ══════════════════════════════════════
import io  # local import kept near usage; io is stdlib, always available

DOC_MODEL = "nvidia/nemotron-3-super-120b-a12b:free"

DOC_TYPES = {"report", "resume", "article", "letter", "proposal", "summary", "essay"}
DOC_FORMATS = {"md", "docx", "pdf"}

DOC_SYSTEM_TMPL = (
    "You are a professional document writer. Write a well-structured {doc_type} on the "
    "given topic, in Markdown. Use a single top-level # heading as the title, ## for "
    "sections, and standard Markdown (bold, lists, etc). No preamble or meta-commentary "
    "— output only the document itself."
)


def _md_to_docx_bytes(markdown_text: str, title: str) -> bytes:
    """Converts simple Markdown (headings, bold, bullet/numbered lists, paragraphs)
    into a .docx using python-docx. Intentionally simple — not a full CommonMark parser."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            p = doc.add_paragraph()
            # Handle **bold** inline segments without a full markdown parser
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _md_to_pdf_bytes(markdown_text: str, title: str) -> bytes:
    """Converts simple Markdown into a .pdf using reportlab Platypus."""
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
    from reportlab.lib.styles import getSampleStyleSheet
    from xml.sax.saxutils import escape

    styles = getSampleStyleSheet()
    story = []

    def inline_bold(text: str) -> str:
        # Convert **bold** to reportlab's <b> markup, escaping everything else first
        escaped = escape(text)
        return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            story.append(Spacer(1, 8))
            continue
        if line.startswith("### "):
            story.append(Paragraph(inline_bold(line[4:]), styles["Heading3"]))
        elif line.startswith("## "):
            story.append(Paragraph(inline_bold(line[3:]), styles["Heading2"]))
        elif line.startswith("# "):
            story.append(Paragraph(inline_bold(line[2:]), styles["Title"]))
        elif re.match(r"^[-*]\s+", line):
            item_text = inline_bold(re.sub(r"^[-*]\s+", "", line))
            story.append(ListFlowable([ListItem(Paragraph(item_text, styles["Normal"]))], bulletType="bullet"))
        else:
            story.append(Paragraph(inline_bold(line), styles["Normal"]))

    buf = io.BytesIO()
    SimpleDocTemplate(buf, pagesize=letter, title=title).build(story)
    return buf.getvalue()


def tool_generate_document(args: dict, api_key: str = None) -> tuple[dict | None, str | None]:
    topic    = (args.get("topic") or "").strip()
    doc_type = (args.get("doc_type") or "report").strip().lower()
    fmt      = (args.get("format") or "md").strip().lower()

    if not topic:
        return None, "Please provide a topic for the document."
    if len(topic) > 500:
        return None, "Topic is too long — keep it under 500 characters."
    if doc_type not in DOC_TYPES:
        return None, f"Unknown document type '{doc_type}'. Choose one of: {', '.join(sorted(DOC_TYPES))}."
    if fmt not in DOC_FORMATS:
        return None, f"Unknown format '{fmt}'. Choose one of: {', '.join(sorted(DOC_FORMATS))}."

    system = DOC_SYSTEM_TMPL.format(doc_type=doc_type)
    markdown_text, err = ai_query_single_model(DOC_MODEL, system, topic, api_key=api_key,
                                                max_tokens=2200, temp=0.5)
    if err:
        return None, err
    if not markdown_text.strip():
        return None, "The model didn't return any content — try rephrasing your topic."

    title_match = re.search(r"^#\s+(.+)$", markdown_text, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else topic[:80]
    safe_stub = re.sub(r"[^a-zA-Z0-9_-]+", "-", title.lower()).strip("-")[:50] or "document"

    result = {
        "markdown": markdown_text,
        "title": title,
        "doc_type": doc_type,
        "requested_format": fmt,
        "actual_format": fmt,
        "fallback_reason": None,
        "filename": f"{safe_stub}.md",
        "file_b64": None,
        "mime": "text/markdown",
    }

    if fmt == "docx":
        try:
            file_bytes = _md_to_docx_bytes(markdown_text, title)
            result.update(
                actual_format="docx", filename=f"{safe_stub}.docx",
                file_b64=base64.b64encode(file_bytes).decode("ascii"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except ImportError:
            result["fallback_reason"] = "python-docx isn't installed on the server — sending Markdown instead."
            result["actual_format"] = "md"
        except Exception as e:
            log.error(f"docx export failed: {e}")
            result["fallback_reason"] = "DOCX export failed unexpectedly — sending Markdown instead."
            result["actual_format"] = "md"

    elif fmt == "pdf":
        try:
            file_bytes = _md_to_pdf_bytes(markdown_text, title)
            result.update(
                actual_format="pdf", filename=f"{safe_stub}.pdf",
                file_b64=base64.b64encode(file_bytes).decode("ascii"),
                mime="application/pdf",
            )
        except ImportError:
            result["fallback_reason"] = "reportlab isn't installed on the server — sending Markdown instead."
            result["actual_format"] = "md"
        except Exception as e:
            log.error(f"pdf export failed: {e}")
            result["fallback_reason"] = "PDF export failed unexpectedly — sending Markdown instead."
            result["actual_format"] = "md"

    if result["file_b64"] is None:
        # md, or graceful fallback from docx/pdf
        result["file_b64"] = base64.b64encode(markdown_text.encode("utf-8")).decode("ascii")

    return result, None


TOOLS = {
    "generate_avatar": {
        "fn": tool_generate_avatar,
        "desc": "Generate a free DiceBear SVG avatar from a seed and style.",
        "needs_key": False,
    },
    "generate_code": {
        "fn": tool_generate_code,
        "desc": "Generate code for a task/language, with static-only syntax verification.",
        "needs_key": True,
    },
    "generate_diagram": {
        "fn": tool_generate_diagram,
        "desc": "Generate a Mermaid.js diagram from a description.",
        "needs_key": True,
    },
    "generate_document": {
        "fn": tool_generate_document,
        "desc": "Generate a document (report/resume/article/etc) exported as md/docx/pdf.",
        "needs_key": True,
    },
    "generate_image": {
        "fn": tool_generate_image,
        "desc": "Generate a custom illustration/scene from a description, via Pollinations (free, no key).",
        "needs_key": False,
    },
}


# ══════════════════════════════════════
# MODULE 6 — Auto-Routing / Tool Calling Integration
# Gives the primary router model (nemotron) OpenAI-style function-calling schemas
# for every registered tool. One request decides: plain reply, or a single tool
# call. Slash commands (Modules 1-5) remain as direct shortcuts and bypass this
# entirely — this only activates for natural-language requests.
# ══════════════════════════════════════
ROUTER_MODEL = "nvidia/nemotron-3-super-120b-a12b:free"

TOOL_SCHEMAS = [
    {
        "type": "function",
        "function": {
            "name": "generate_avatar",
            "description": "Generate a free DiceBear SVG avatar image from a seed (e.g. a name) and a visual style. Use when the user asks for an avatar, profile picture, or icon to be generated/created.",
            "parameters": {
                "type": "object",
                "properties": {
                    "seed": {"type": "string", "description": "Text to derive the avatar from, e.g. a name."},
                    "style": {"type": "string", "enum": sorted(DICEBEAR_STYLES), "description": "DiceBear avatar style."},
                },
                "required": ["seed"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_code",
            "description": "Generate source code for a described task in a given programming language, with a static (non-executing) syntax check. Use when the user asks to write, create, or generate code, a script, or a function.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "Description of the code to generate."},
                    "language": {"type": "string", "description": "Target programming language, e.g. python, javascript."},
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_diagram",
            "description": "Generate a Mermaid.js diagram (flowchart, sequence diagram, class diagram, etc) from a description. Use when the user asks for a diagram, flowchart, chart of a process, architecture visual, or similar.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "Description of the diagram/process/system to visualize."},
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_document",
            "description": "Generate a written document (report, resume, article, letter, proposal, summary, or essay) on a topic, exported as Markdown, DOCX, or PDF. Use when the user asks to write/draft/create a document, report, resume, article, or similar deliverable, especially if they want it downloadable.",
            "parameters": {
                "type": "object",
                "properties": {
                    "topic": {"type": "string", "description": "The subject/topic of the document."},
                    "doc_type": {"type": "string", "enum": sorted(DOC_TYPES), "description": "Kind of document."},
                    "format": {"type": "string", "enum": sorted(DOC_FORMATS), "description": "Export file format."},
                },
                "required": ["topic"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_image",
            "description": "Generate a custom illustration, picture, or scene from a text description — e.g. 'a cat eating ice cream', 'a sunset over mountains', 'a robot playing guitar'. Use this for any request to draw, create, generate, paint, or make a PICTURE or IMAGE of something. Distinct from generate_avatar, which only makes small stylized profile-icon avatars — use generate_image for anything more detailed or scene-like. If the user implies a shape (portrait, landscape, passport photo, square, wallpaper, banner, 16:9, 9:16, etc.), keep that wording in the prompt you pass — the renderer picks image dimensions from it.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "Description of the image/scene to generate."},
                },
                "required": ["prompt"],
            },
        },
    },
]


def route_intent(user_msg: str, history: list = None, api_key: str = None):
    """Single function-calling request to the router model. Returns a dict:
    {"type": "tool", "tool": name, "args": {...}} if the model wants a tool,
    {"type": "text", "text": "..."} if it answered directly,
    or {"type": "none"} on any failure — callers should fall through to the
    normal ai_query path in that case, so auto-routing never blocks a reply."""
    key = (api_key or AI_API_KEY or "").strip()
    if not key:
        return {"type": "none"}

    messages = [{
        "role": "system",
        "content": (
            f"You are {BOT_NAME}'s routing layer. Decide whether the user's message requires "
            "calling one of the available tools (avatar/code/diagram/document generation), or "
            "is a normal conversational message you should just answer directly. Only call a "
            "tool when the user is clearly asking to generate/create/write/draw one of those "
            "specific artifacts. For everything else — questions, chat, opinions — respond "
            "directly with no tool call."
        ),
    }]
    if history:
        for m in history[-8:]:
            if m.get("role") in ("user", "assistant") and m.get("content"):
                messages.append({"role": m["role"], "content": m["content"]})
    messages.append({"role": "user", "content": user_msg})

    headers = {
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://eka-dev1.onrender.com",
        "X-Title": f"{BOT_NAME} Router",
    }
    body = {
        "model": ROUTER_MODEL,
        "messages": messages,
        "tools": TOOL_SCHEMAS,
        "tool_choice": "auto",
        "max_tokens": 700,
        "temperature": 0.2,
    }

    try:
        resp = requests.post(AI_API_URL, headers=headers, json=body, timeout=30)
        if resp.status_code != 200:
            log.warning(f"router: HTTP {resp.status_code}: {resp.text[:200]}")
            return {"type": "none"}

        choice = (resp.json().get("choices") or [{}])[0]
        message = choice.get("message", {})
        tool_calls = message.get("tool_calls") or []

        if tool_calls:
            call = tool_calls[0]
            fn = call.get("function", {})
            name = fn.get("name")
            if name not in TOOLS:
                return {"type": "none"}
            try:
                args = json.loads(fn.get("arguments") or "{}")
            except (json.JSONDecodeError, TypeError):
                args = {}
            return {"type": "tool", "tool": name, "args": args}

        text = (message.get("content") or "").strip()
        if text:
            return {"type": "text", "text": clean(text)}

        return {"type": "none"}
    except requests.exceptions.Timeout:
        log.warning("router: timeout")
        return {"type": "none"}
    except Exception as e:
        log.error(f"router error: {e}")
        return {"type": "none"}


# ══════════════════════════════════════
# RESPONSE CLEANING
# ══════════════════════════════════════
def clean(text: str) -> str:
    if not text:
        return ""
    # Strip internal <think> blocks some models emit
    text = re.sub(r"<think(?:ing)?>.*?</think(?:ing)?>", "", text, flags=re.DOTALL | re.IGNORECASE)
    # Strip self-labelling prefix
    text = re.sub(r"^(EKA\s*:\s*|Eka\s*:\s*|Assistant\s*:\s*)", "", text, flags=re.IGNORECASE)
    # Remove stray XML tags
    text = re.sub(r"</?[a-zA-Z_][^>]{0,50}>", "", text)
    # Collapse 3+ blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ══════════════════════════════════════
# AI CORE
# ══════════════════════════════════════
def ai_query(user_input: str, history: list = None, system: str = None, image_data_url: str = None, api_key: str = None) -> str:
    key = (api_key or AI_API_KEY or "").strip()
    if not key:
        return "I don't have an AI key configured yet, so I can't reply. Please add your own free OpenRouter key: Menu > Settings > Your API Key (BYOK)."

    messages = [{"role": "system", "content": system or SYS_BASE}]

    if history:
        for m in history[-16:]:
            if m.get("role") in ("user", "assistant") and m.get("content"):
                messages.append({"role": m["role"], "content": m["content"]})

    if image_data_url:
        # Multimodal content — only vision-capable models in the waterfall will be tried below.
        messages.append({"role": "user", "content": [
            {"type": "text", "text": user_input or "Please describe this image."},
            {"type": "image_url", "image_url": {"url": image_data_url}},
        ]})
    else:
        messages.append({"role": "user", "content": user_input})

    headers = {
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://eka-dev1.onrender.com",
        "X-Title": f"{BOT_NAME} AI",
    }

    models_to_try = [m for m in MODELS if not image_data_url or m.get("vision")]
    if image_data_url and not models_to_try:
        return "None of the configured models support image input right now."

    for model in models_to_try:
        try:
            t0   = time.time()
            body = {"model": model["id"], "messages": messages,
                    "max_tokens": model["max_tokens"], "temperature": model["temp"]}
            resp = requests.post(AI_API_URL, headers=headers, json=body, timeout=35)
            log.info(f"  {model['id']} → {resp.status_code} ({round(time.time()-t0,2)}s)")

            if resp.status_code == 200:
                data   = resp.json()
                choice = (data.get("choices") or [{}])[0]
                text   = (choice.get("message") or {}).get("content", "").strip()
                if text:
                    return clean(text)
                log.warning(f"  Empty reply from {model['id']} — raw: {resp.text[:500]}")

            elif resp.status_code == 429:
                log.warning(f"  Rate-limited on {model['id']}, trying next… body: {resp.text[:500]}")
                time.sleep(2.0)

            elif 400 <= resp.status_code < 500:
                log.warning(f"  Client error {resp.status_code} on {model['id']}, skipping — body: {resp.text[:500]}")

            else:
                log.warning(f"  Server error {resp.status_code} on {model['id']} — body: {resp.text[:500]}")

        except requests.exceptions.Timeout:
            log.warning(f"  Timeout on {model['id']}")
        except Exception as e:
            log.error(f"  Error on {model['id']}: {e}")

    return "All AI models are temporarily unavailable. Please try again shortly."


def ai_query_single_model(model_id: str, system: str, user_input: str, api_key: str = None,
                           max_tokens: int = 1500, temp: float = 0.3) -> tuple[str | None, str | None]:
    """Like ai_query but targets one specific model (no waterfall) — used by tools
    that need a particular model's strengths (e.g. code generation). Falls back to
    the first model in MODELS if the requested one errors. Returns (text, error)."""
    key = (api_key or AI_API_KEY or "").strip()
    if not key:
        return None, "No AI key configured. Add your own OpenRouter key: Menu > Settings > Your API Key (BYOK), or contact the admin."

    headers = {
        "Authorization": f"Bearer {key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://eka-dev1.onrender.com",
        "X-Title": f"{BOT_NAME} AI",
    }
    messages = [{"role": "system", "content": system}, {"role": "user", "content": user_input}]

    candidates = [model_id] + [m["id"] for m in MODELS if m["id"] != model_id]
    for mid in candidates:
        try:
            t0   = time.time()
            body = {"model": mid, "messages": messages, "max_tokens": max_tokens, "temperature": temp}
            resp = requests.post(AI_API_URL, headers=headers, json=body, timeout=45)
            log.info(f"  {mid} → {resp.status_code} ({round(time.time()-t0,2)}s)")

            if resp.status_code == 200:
                data   = resp.json()
                choice = (data.get("choices") or [{}])[0]
                text   = (choice.get("message") or {}).get("content", "").strip()
                if text:
                    return clean(text), None
                log.warning(f"  Empty reply from {mid}")
            elif resp.status_code == 429:
                log.warning(f"  Rate-limited on {mid}, trying next…")
                time.sleep(1.5)
            else:
                log.warning(f"  Error {resp.status_code} on {mid}: {resp.text[:300]}")
        except requests.exceptions.Timeout:
            log.warning(f"  Timeout on {mid}")
        except Exception as e:
            log.error(f"  Error on {mid}: {e}")

    return None, "All AI models are temporarily unavailable. Please try again shortly."


# ══════════════════════════════════════
# QUICK REPLIES (no AI cost)
# ══════════════════════════════════════
def quick_reply(text: str) -> str | None:
    t = text.lower().strip().rstrip("?!.,")
    greetings = {"hi","hello","hey","namaste","namaskar","hola","yo","hii","hai","hyy","good morning","good evening","good night","good afternoon"}
    if t in greetings:
        return f"Hey! 👋 I'm **{BOT_NAME}**, your AI assistant, built in India 🇮🇳. What can I help you with?"

    identity = re.search(r"\b(who are you|your name|what are you|introduce yourself|aap kaun|tumhara naam|kaun ho)\b", t)
    if identity:
        return f"I'm **{BOT_NAME}** — an AI assistant built by **{DEV_NAME}** in India 🇮🇳. I can help with questions, code, writing, analysis, and more. Ask away!"

    return None


# ══════════════════════════════════════
# MODULE 4 — Auth: SQLite user store
# stdlib sqlite3, no new deps. Passwords hashed with PBKDF2-HMAC-SHA256 + per-user salt.
# ══════════════════════════════════════
DB_PATH = os.getenv("AUTH_DB_PATH", os.path.join(os.path.dirname(os.path.abspath(__file__)), "eka_users.db"))


def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with get_db() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                username     TEXT UNIQUE NOT NULL,
                display_name TEXT NOT NULL,
                pw_hash      TEXT NOT NULL,
                pw_salt      TEXT NOT NULL,
                created_at   TEXT NOT NULL,
                avatar       TEXT,
                about        TEXT
            )
        """)
        # Migration guard: existing DBs (like eka_users.db already in this repo)
        # were created before avatar/about existed — add them if missing so we
        # don't need people to delete their user database to get this update.
        existing_cols = {row["name"] for row in conn.execute("PRAGMA table_info(users)")}
        if "avatar" not in existing_cols:
            conn.execute("ALTER TABLE users ADD COLUMN avatar TEXT")
        if "about" not in existing_cols:
            conn.execute("ALTER TABLE users ADD COLUMN about TEXT")


init_db()


def hash_password(password: str, salt: str = None) -> tuple[str, str]:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 260_000)
    return digest.hex(), salt


def create_user(username: str, display_name: str, password: str) -> tuple[dict | None, str | None]:
    username = username.strip().lower()
    display_name = display_name.strip()

    if not re.match(r"^[a-zA-Z0-9_.]{3,32}$", username):
        return None, "Username must be 3-32 characters: letters, numbers, underscore, or period only."
    if not display_name:
        return None, "Please enter a display name."
    if len(password) < 4:
        return None, "Password must be at least 4 characters."

    pw_hash, pw_salt = hash_password(password)
    try:
        with get_db() as conn:
            conn.execute(
                "INSERT INTO users (username, display_name, pw_hash, pw_salt, created_at) VALUES (?, ?, ?, ?, ?)",
                (username, display_name, pw_hash, pw_salt, datetime.utcnow().isoformat()),
            )
            row = conn.execute("SELECT id, username, display_name, avatar, about FROM users WHERE username = ?", (username,)).fetchone()
        return dict(row), None
    except sqlite3.IntegrityError:
        return None, "That username is already taken."


def verify_user(username: str, password: str) -> tuple[dict | None, str | None]:
    username = username.strip().lower()
    with get_db() as conn:
        row = conn.execute("SELECT * FROM users WHERE username = ?", (username,)).fetchone()

    if not row:
        return None, "Invalid username or password."

    check_hash, _ = hash_password(password, row["pw_salt"])
    if not secrets.compare_digest(check_hash, row["pw_hash"]):
        return None, "Invalid username or password."

    return {"id": row["id"], "username": row["username"], "display_name": row["display_name"],
            "avatar": row["avatar"], "about": row["about"]}, None


def get_user_by_id(user_id: int) -> dict | None:
    with get_db() as conn:
        row = conn.execute("SELECT id, username, display_name, avatar, about FROM users WHERE id = ?", (user_id,)).fetchone()
    return dict(row) if row else None


def update_user_profile(user_id: int, display_name: str | None = None, about: str | None = None) -> tuple[dict | None, str | None]:
    display_name = (display_name or "").strip()
    if not display_name:
        return None, "Display name can't be empty."
    if len(display_name) > 32:
        return None, "Display name is too long (max 32 characters)."
    about = (about or "").strip()[:60]
    with get_db() as conn:
        conn.execute("UPDATE users SET display_name = ?, about = ? WHERE id = ?", (display_name, about, user_id))
    return get_user_by_id(user_id), None


def update_user_avatar(user_id: int, data_url: str | None) -> tuple[dict | None, str | None]:
    # data_url is either a "data:image/...;base64,..." string or None to remove the photo
    if data_url:
        if not data_url.startswith("data:image/"):
            return None, "Invalid image data."
        if len(data_url) > 1_800_000:  # ~1.3MB decoded — plenty for a profile photo, keeps DB rows small
            return None, "Image is too large. Please choose a smaller photo."
    with get_db() as conn:
        conn.execute("UPDATE users SET avatar = ? WHERE id = ?", (data_url, user_id))
    return get_user_by_id(user_id), None


# ══════════════════════════════════════
# BYOK — Bring Your Own OpenRouter Key
# Frontend sends the user's key (from localStorage) via header on each request.
# We accept either header name for flexibility; falls back to server env key.
# ══════════════════════════════════════
def get_user_api_key() -> str | None:
    key = request.headers.get("Authorization", "").strip()
    if key.lower().startswith("bearer "):
        key = key[7:].strip()
    if not key:
        key = (request.headers.get("X-API-Key") or "").strip()
    return key or None


# ══════════════════════════════════════
# AUTH GUARD
# The chat UI can be served from this Flask app OR from a separate static host
# (GitHub Pages) that calls these APIs cross-origin with credentials — so login
# state can't just be checked server-side when rendering the page in that case.
# require_login() covers the same-origin path; /api/me covers the cross-origin one.
# ══════════════════════════════════════
def require_login():
    return "user_id" in session


# ══════════════════════════════════════
# ROUTES
# ══════════════════════════════════════
@app.route("/")
def index():
    # Gate the chat page itself: no session → send to /login first.
    # (When the chat UI is hosted separately on GitHub Pages, that copy calls
    # GET /api/me on load and redirects client-side instead — see script.js.)
    if not require_login():
        return redirect(url_for("login_page"))
    return render_template("index.html", bot_name=BOT_NAME)


@app.route("/login")
def login_page():
    if session.get("user_id"):
        return redirect(url_for("index"))
    return render_template("login.html", bot_name=BOT_NAME)


@app.route("/signup")
def signup_page():
    if session.get("user_id"):
        return redirect(url_for("index"))
    return render_template("signup.html", bot_name=BOT_NAME)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login_page"))


@app.route("/api/logout", methods=["POST"])
def api_logout():
    # JSON/fetch-friendly logout for the cross-origin (GitHub Pages) frontend,
    # which can't follow a same-origin redirect the way the Flask-served page can.
    session.clear()
    return jsonify({"ok": True, "redirect": url_for("login_page", _external=True)})


@app.route("/api/me")
def api_me():
    # Cross-origin login check the standalone frontend calls on page load.
    if not require_login():
        return jsonify({"error": "Not logged in"}), 401
    user = get_user_by_id(session["user_id"])
    if not user:
        session.clear()
        return jsonify({"error": "Not logged in"}), 401
    return jsonify({"user": user})


@app.route("/api/profile", methods=["POST"])
def api_update_profile():
    if not require_login():
        return jsonify({"error": "Not logged in"}), 401
    payload = request.get_json(silent=True) or {}
    user, err = update_user_profile(
        session["user_id"],
        display_name=payload.get("display_name"),
        about=payload.get("about"),
    )
    if err:
        return jsonify({"error": err}), 400
    session["display_name"] = user["display_name"]
    log.info(f"→ profile updated: {user['username']}")
    return jsonify({"user": user})


@app.route("/api/profile/avatar", methods=["POST"])
def api_update_avatar():
    if not require_login():
        return jsonify({"error": "Not logged in"}), 401
    payload = request.get_json(silent=True) or {}
    user, err = update_user_avatar(session["user_id"], payload.get("avatar"))
    if err:
        return jsonify({"error": err}), 400
    log.info(f"→ avatar updated: {user['username']}")
    return jsonify({"user": user})


@app.route("/api/signup", methods=["POST"])
def api_signup():
    payload = request.get_json(silent=True) or {}
    if (payload.get("website") or "").strip():
        # Honeypot field from the signup form — real users never fill it.
        return jsonify({"error": "Sign up failed"}), 400
    user, err = create_user(
        username=payload.get("username", ""),
        display_name=payload.get("display_name", ""),
        password=payload.get("password", ""),
    )
    if err:
        return jsonify({"error": err}), 400

    session.clear()
    session["user_id"] = user["id"]
    session["username"] = user["username"]
    session["display_name"] = user["display_name"]
    log.info(f"→ signup: {user['username']}")
    return jsonify({"redirect": FRONTEND_URL, "user": user})


@app.route("/api/login", methods=["POST"])
def api_login():
    payload = request.get_json(silent=True) or {}
    user, err = verify_user(payload.get("username", ""), payload.get("password", ""))
    if err:
        return jsonify({"error": err}), 401

    session.clear()
    session["user_id"] = user["id"]
    session["username"] = user["username"]
    session["display_name"] = user["display_name"]
    log.info(f"→ login: {user['username']}")
    return jsonify({"redirect": FRONTEND_URL, "user": user})


@app.route("/api/chat", methods=["POST"])
def chat():
    payload  = request.get_json(silent=True) or {}
    user_msg = (payload.get("message") or "").strip()
    history  = payload.get("history", [])
    use_web  = payload.get("wiki", False)
    image    = payload.get("image")  # optional base64 data-URL of an attached photo
    user_key = get_user_api_key()    # BYOK — user's own OpenRouter key, if provided

    if not user_msg and not image:
        return jsonify({"reply": "Your message seems empty. What would you like to ask?", "source": "system"})

    log.info(f"→ {user_msg[:80]}{' [+image]' if image else ''}{' [BYOK]' if user_key else ''}")

    # If the person is logged in, let the AI know who it's talking to — pulled
    # from the server-side session (not the request body), so it can't be spoofed
    # by editing frontend JS, and works without the frontend having to pass it.
    user_name_line = ""
    if require_login():
        display_name = (session.get("display_name") or "").strip()
        if display_name:
            user_name_line = f"\nThe person you're talking to is named {display_name} — address them by name when it feels natural, don't force it into every reply."

    # Image path — route straight to a vision-capable model, skip quick-replies/web-search
    if image:
        system = (SYS_BASE + user_name_line) if user_name_line else None
        reply = ai_query(user_msg, history=history, system=system, image_data_url=image, api_key=user_key)
        log.info(f"← ai+vision: {reply[:60]}")
        return jsonify({"reply": reply, "source": "ai"})

    # Quick path
    quick = quick_reply(user_msg)
    if quick:
        return jsonify({"reply": quick, "source": "system"})

    # ── Auto-routing (Module 6) — let the router model decide if a tool fits ──
    # Failures fall through silently to the normal web-search/AI path below;
    # auto-routing never blocks a reply.
    routed = route_intent(user_msg, history=history, api_key=user_key)

    if routed["type"] == "tool":
        tool_name = routed["tool"]
        tool = TOOLS[tool_name]
        if tool.get("needs_key") and not (user_key or AI_API_KEY):
            log.warning(f"router picked {tool_name} but no key available — falling back to text")
        else:
            data, err = tool["fn"](routed["args"], api_key=user_key)
            log.info(f"← router→tool:{tool_name} {'ok' if not err else 'error: ' + err}")
            if err:
                return jsonify({"reply": err, "source": "system"})
            return jsonify({"reply": None, "source": "tool", "tool": tool_name, "tool_data": data})

    if routed["type"] == "text":
        log.info(f"← router (direct reply): {routed['text'][:60]}")
        return jsonify({"reply": routed["text"], "source": "ai"})

    # Web search path
    if use_web:
        content, src = web_search(user_msg)
        if content:
            system = SYS_WEB.replace("{web_content}", content) + user_name_line
            reply  = ai_query(user_msg, history=history, system=system, api_key=user_key)
            log.info(f"← web+ai [{src}]: {reply[:60]}")
            return jsonify({"reply": reply, "source": "web+ai", "web_source": src})

    # Standard AI
    system = (SYS_BASE + user_name_line) if user_name_line else None
    reply = ai_query(user_msg, history=history, system=system, api_key=user_key)
    log.info(f"← ai: {reply[:60]}")
    return jsonify({"reply": reply, "source": "ai"})


@app.route("/api/image", methods=["POST"])
def image():
    payload = request.get_json(silent=True) or {}
    prompt  = (payload.get("prompt") or "").strip()

    if not prompt:
        return jsonify({"error": "Describe what you'd like me to draw."}), 400
    if len(prompt) > 600:
        return jsonify({"error": "That prompt is a bit long — try trimming it."}), 400

    log.info(f"→ image: {prompt[:80]}")
    data_url, err = generate_image(prompt)
    if err:
        log.warning(f"← image failed: {err}")
        return jsonify({"error": err}), 502

    log.info("← image: ok")
    return jsonify({"image": data_url, "prompt": prompt, "source": "pollinations"})


@app.route("/api/tool/<tool_name>", methods=["POST"])
def run_tool(tool_name):
    tool = TOOLS.get(tool_name)
    if not tool:
        return jsonify({"ok": False, "tool": tool_name, "error": "Unknown tool."}), 404

    args = request.get_json(silent=True) or {}
    user_key = get_user_api_key()  # BYOK — forwarded to every tool; LLM-backed tools use it, others ignore it

    if tool.get("needs_key") and not (user_key or AI_API_KEY):
        return jsonify({
            "ok": False, "tool": tool_name,
            "error": "This tool needs an AI key. Add your own OpenRouter key: Menu > Settings > Your API Key (BYOK), or contact the admin."
        }), 400

    log.info(f"→ tool:{tool_name} args={args}{' [BYOK]' if user_key else ''}")
    data, err = tool["fn"](args, api_key=user_key)

    if err:
        log.warning(f"← tool:{tool_name} error: {err}")
        return jsonify({"ok": False, "tool": tool_name, "error": err}), 400

    log.info(f"← tool:{tool_name} ok")
    return jsonify({"ok": True, "tool": tool_name, "data": data})


@app.route("/api/health")
def health():
    return jsonify({"status": "ok", "bot": BOT_NAME,
                    "models": [m["id"] for m in MODELS],
                    "has_server_key": bool((AI_API_KEY or "").strip()),
                    "time": datetime.now().isoformat()})


# Serve the static 404.html (shared with GitHub Pages) for unmatched routes
# on the Flask side too, so both deployment targets show the same page.
@app.errorhandler(404)
def not_found(e):
    return send_from_directory(app.root_path, "404.html"), 404


if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    log.info(f"Starting {BOT_NAME} AI on :{port}")
    app.run(debug=os.getenv("DEBUG","true").lower()=="true", host="0.0.0.0", port=port)
